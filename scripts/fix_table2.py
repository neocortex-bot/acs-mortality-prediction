import docx
from docx.shared import Pt
from lxml import etree
from copy import deepcopy
from docx.oxml.ns import qn, nsdecls

doc = docx.Document('TESIS_FINAL.docx')
t = doc.tables[0]
print(f'Starting rows: {len(t.rows)}')

def new_row(table, ref_row_idx):
    from docx.oxml import OxmlElement
    ref_row = table.rows[ref_row_idx]
    ref_tr = ref_row._tr
    ref_tcs = list(ref_tr.iterchildren(qn('w:tc')))
    if not ref_tcs:
        return None
    ref_tc = ref_tcs[0]
    ref_tcPr = ref_tc.find(qn('w:tcPr'))
    
    new_tr = OxmlElement('w:tr')
    
    # Copy trPr
    trPr = ref_tr.find(qn('w:trPr'))
    if trPr is not None:
        new_tr.append(deepcopy(trPr))
    
    for i in range(4):
        tc = OxmlElement('w:tc')
        tcPr = OxmlElement('w:tcPr')
        tc.append(tcPr)
        if ref_tcPr is not None:
            for child in ref_tcPr:
                if child.tag not in [qn('w:gridSpan'), qn('w:vMerge')]:
                    tcPr.append(deepcopy(child))
        p = OxmlElement('w:p')
        tc.append(p)
        r = OxmlElement('w:r')
        p.append(r)
        t_elem = OxmlElement('w:t')
        t_elem.text = ''
        r.append(t_elem)
        new_tr.append(tc)
    
    ref_tr.addnext(new_tr)
    
    cells = []
    for tc in new_tr.iterchildren(qn('w:tc')):
        p_elem = tc.find(qn('w:p'))
        new_para = docx.text.paragraph.Paragraph(p_elem, tc)
        cells.append(docx.table._Cell(tc, new_para))
    return cells

def sc(cells, idx, text, size=Pt(10)):
    if idx < len(cells):
        cell = cells[idx]
        p = cell.paragraphs[0]
        p.clear()
        r = p.add_run(text)
        r.font.size = size

offset = 0
errors = []

def safe_insert(ref, label, *pairs):
    global offset
    c = new_row(t, ref)
    if c is None:
        errors.append(f'{label}: new_row returned None')
        return
    for i in range(0, len(pairs), 2):
        col = pairs[i]
        txt = pairs[i+1]
        sc(c, col, txt)
    offset += 1
    print(f'  OK R{ref} {label}')

safe_insert(26, 'NLR', 0, 'NLR (Neutrophil-to-Lymphocyte Ratio)', 1, 'Rasio neutrofil terhadap limfosit dari darah tepi, nilai pada 0-24 jam pertama', 2, 'Kontinu', 3, 'NLR = hitung jenis neutrofil absolut / hitung jenis limfosit absolut')
safe_insert(26 + offset, 'SII', 0, 'SII (Systemic Immune-Inflammation Index)', 1, 'Indeks imun-inflamasi sistemik, dihitung dari nilai laboratorium 0-24 jam pertama', 2, 'Kontinu', 3, 'SII = (hitung jenis neutrofil absolut × hitung jenis trombosit absolut) / hitung jenis limfosit absolut')
safe_insert(26 + offset, 'APTT', 0, 'APTT (Activated Partial Thromboplastin Time)', 1, 'Waktu tromboplastin parsial teraktivasi, nilai pada 0-24 jam pertama', 2, 'Kontinu', 3, 'detik, nilai numerik')
safe_insert(23 + offset, 'Kreatinin', 0, 'Kreatinin', 1, 'Kadar kreatinin serum, nilai pada 0-24 jam pertama', 2, 'Kontinu', 3, 'mg/dL, nilai numerik')
safe_insert(12 + offset, 'RR', 0, 'Laju Napas (Respiratory Rate)', 1, 'Frekuensi napas saat admisi, nilai tertinggi (maksimum) dalam 0-24 jam pertama', 2, 'Kontinu', 3, 'x/menit, nilai numerik')
safe_insert(30 + offset, 'LVOT VTI', 0, 'LVOT VTI (Velocity Time Integral)', 1, 'Velocity time integral dari left ventricular outflow tract, pulsed-wave Doppler apikal 5-ruang', 2, 'Kontinu', 3, 'cm, nilai numerik')
safe_insert(38 + offset, 'Dislipidemia', 0, 'Dislipidemia', 1, 'Riwayat dislipidemia atau penggunaan obat penurun lipid sebelum admisi', 2, 'Biner', 3, 'Ya=1 jika di problem list atau penggunaan statin/fibrat kronik sebelum admisi; Tidak=0')
safe_insert(38 + offset, 'PGK', 0, 'Penyakit Ginjal Kronik', 1, 'Riwayat penyakit ginjal kronik atau eGFR <60 mL/menit/1,73 m2 yang menetap sebelum admisi', 2, 'Biner', 3, 'Ya=1 jika di problem list atau eGFR kronik <60; Tidak=0')
safe_insert(15 + offset, 'Elevasi ST', 0, 'Elevasi ST', 1, 'Terdapat elevasi segmen ST pada EKG awal sesuai definisi STEMI universal', 2, 'Biner', 3, 'Ya=1 jika elevasi ST >=1 mm pada >=2 sadapan ekstremitas atau >=2 mm pada >=2 sadapan prekordial; Tidak=0')
safe_insert(15 + offset, 'Gel Q', 0, 'Gelombang Q patologis/OMI', 1, 'Terdapat gelombang Q patologis atau gambaran old myocardial infarction pada EKG awal', 2, 'Biner', 3, 'Ya=1 jika durasi Q >=0,04 detik dan/atau kedalaman Q >=25% tinggi R pada >=2 sadapan; Tidak=0')
safe_insert(13 + offset, 'Kardiomegali', 0, 'Kardiomegali', 1, 'Pembesaran ukuran jantung berdasarkan foto toraks saat admisi', 2, 'Biner', 3, 'Ya=1 jika CTR (cardiothoracic ratio) >0,5 pada foto toraks PA; Tidak=0')
safe_insert(13 + offset, 'Edema Paru', 0, 'Edema Paru', 1, 'Tanda edema paru berdasarkan foto toraks atau pemeriksaan fisik saat admisi', 2, 'Biner', 3, 'Ya=1 jika terdapat gambaran edema paru pada foto toraks atau ronkhi basah pada auskultasi paru; Tidak=0')

if errors:
    print(f'\nERRORS: {errors}')
else:
    print(f'\nAll {offset} rows inserted successfully')
    
doc.save('TESIS_FINAL.docx')
print(f'Total rows now: {len(doc.tables[0].rows)}')
