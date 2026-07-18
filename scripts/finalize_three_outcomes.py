#!/usr/bin/env python3
"""Synchronize the three-outcome thesis narrative, table, and notebook section."""

import json
from pathlib import Path

import nbformat
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "TESIS_FINAL.docx"
SOURCE_NB = ROOT / "gm-acs-mortality-prediction.ipynb"


def set_paragraph_text(paragraph, text):
    """Replace text while retaining paragraph properties and first-run formatting."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def shade(cell, fill="D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def finalize_docx():
    doc = Document(DOCX)

    # Resolve a legacy duplicate figure number and its reporting-map reference.
    set_paragraph_text(doc.paragraphs[136], "Gambar 2.3 Diagram alur seleksi pasien penelitian")
    set_paragraph_text(
        doc.paragraphs[161],
        "Kurva ROC pada Gambar 3.1 menunjukkan performa diskriminasi model pada berbagai ambang "
        "klasifikasi. AUC 0,819 menunjukkan kemampuan yang baik untuk membedakan pasien yang "
        "mengalami mortalitas dari pasien yang hidup selama perawatan."
    )
    for table in doc.tables:
        if table.cell(0, 0).text == "Item" and "Lokasi dan status pelaporan" in table.cell(0, 1).text:
            for row in table.rows:
                if row.cells[0].text.startswith("20. Alur partisipan"):
                    row.cells[1].text = "Bagian 3.1 dan Gambar 2.3 memuat seluruh tahap seleksi dan jumlah pasien."

    # Correct the residual legacy prevalence and provide class-imbalance context.
    set_paragraph_text(
        doc.paragraphs[219],
        "Model menunjukkan AUC 0,819 untuk mortalitas, 0,769 untuk luaran komposit, dan "
        "0,747 untuk syok kardiogenik baru. Dari 1.524 subjek, terdapat 115 kematian "
        "(7,5%), 171 kejadian syok kardiogenik baru (11,2%), dan 197 kejadian komposit "
        "(12,9%). AUPRC syok kardiogenik adalah 0,500 dibandingkan prevalensi 0,112, "
        "sedangkan AUPRC komposit adalah 0,635 dibandingkan prevalensi 0,129. Perbandingan "
        "deskriptif ini menunjukkan pemeringkatan kelas minoritas yang lebih baik daripada "
        "baseline tanpa keterampilan. Analisis luaran sekunder bersifat eksploratif dan tidak "
        "mengubah fokus utama penelitian pada mortalitas in-hospital."
    )
    set_paragraph_text(
        doc.paragraphs[268],
        "Performa model untuk syok kardiogenik baru (AUC 0,747) lebih rendah daripada "
        "mortalitas (AUC 0,819). Beberapa faktor dapat menjelaskan perbedaan ini. Pertama, "
        "definisi dan pencatatan syok kardiogenik dalam rekam medis mungkin tidak seragam "
        "karena desain retrospektif. Kedua, sebagian pasien yang diklasifikasikan mengalami "
        "syok baru mungkin telah berada pada fase awal saat admisi, tetapi belum terdokumentasi "
        "secara lengkap. Ketiga, 171 kejadian (11,2%) tetap menyediakan informasi kelas positif "
        "yang lebih terbatas daripada kelas negatif. Berdasarkan kepentingan fitur, LVOT VTI "
        "(Gini 0,132), SBP (0,128), dan APTT (0,083) adalah tiga prediktor terkuat."
    )
    set_paragraph_text(
        doc.paragraphs[271],
        "Luaran komposit mortalitas dan/atau syok kardiogenik menunjukkan AUC 0,769, di antara "
        "performa mortalitas (0,819) dan syok kardiogenik (0,747). Analisis ini bersifat "
        "eksploratif karena menggabungkan komponen dengan karakteristik prediktif dan implikasi "
        "klinis yang berbeda. Berdasarkan kepentingan fitur, LVOT VTI (Gini 0,134), SBP "
        "(0,117), dan eGFR (0,113) adalah tiga prediktor terkuat untuk luaran komposit."
    )
    set_paragraph_text(
        doc.paragraphs[272],
        "Dari perspektif klinis, luaran komposit meningkatkan jumlah kejadian dan kekuatan "
        "statistik, tetapi interpretasinya lebih kompleks karena setiap komponen dapat memiliki "
        "implikasi tata laksana yang berbeda. Dalam konteks IGD, prediksi mortalitas tetap menjadi "
        "fokus utama karena berhubungan langsung dengan kebutuhan penilaian dan intervensi segera."
    )

    # Add the requested comparison table immediately after the composite discussion.
    headers = ["Luaran", "AUC", "±SD", "IK 95%", "Kepentingan fitur (3 tertinggi)"]
    rows = [
        ["Mortalitas", "0,819", "0,007", "0,805-0,833", "eGFR, ureum, LVOT VTI"],
        ["Syok kardiogenik baru", "0,747", "0,005", "0,736-0,757", "LVOT VTI, SBP, APTT"],
        ["Komposit", "0,769", "0,004", "0,761-0,777", "LVOT VTI, SBP, eGFR"],
    ]
    existing = [t for t in doc.tables if t.cell(0, 0).text == "Luaran" and "±SD" in [c.text for c in t.rows[0].cells]]
    if not existing:
        caption = doc.add_paragraph("Tabel 4.1 Perbandingan Performa Tiga Luaran")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        for cell, value in zip(table.rows[0].cells, headers):
            cell.text = value
            shade(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for run in cell.paragraphs[0].runs:
                run.bold = True
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        repeat = OxmlElement("w:tblHeader")
        repeat.set(qn("w:val"), "true")
        tr_pr.append(repeat)
        for values in rows:
            cells = table.add_row().cells
            for cell, value in zip(cells, values):
                cell.text = value
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        anchor = doc.paragraphs[272]._p
        anchor.addnext(caption._p)
        caption._p.addnext(table._tbl)

    # Enforce the requested punctuation rule across prose and tables.
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.text = run.text.replace("—", ",").replace("–", " sampai ")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = run.text.replace("—", ",").replace("–", " sampai ")

    doc.save(DOCX)


def finalize_notebook():
    nb = nbformat.read(SOURCE_NB, as_version=4)
    # Remove any earlier injected copy, making the operation idempotent.
    start = next((i for i, c in enumerate(nb.cells) if c.cell_type == "markdown" and c.source.startswith("## 7. Secondary outcomes")), None)
    if start is not None:
        del nb.cells[start:start + 3]

    insertion = next(
        i for i, c in enumerate(nb.cells)
        if c.cell_type == "markdown" and (
            c.source.startswith("## 7. Reproducibility record")
            or c.source.startswith("## 10. Reproducibility record")
        )
    )
    nb.cells[insertion].source = nb.cells[insertion].source.replace("## 7.", "## 10.", 1)
    for cell in nb.cells[insertion + 1:]:
        if cell.cell_type == "markdown" and (
            cell.source.startswith("## 8. Final synchronized metrics")
            or cell.source.startswith("## 11. Final synchronized metrics")
        ):
            cell.source = cell.source.replace("## 8.", "## 11.", 1)

    code = """with open('three_outcomes_results.json') as f:
    three = json.load(f)

outcome_rows = [
    ('Mortality', three['mortality'], 115, 7.5),
    ('De novo shock', three['shock'], 171, 11.2),
    ('Composite', three['composite'], 197, 12.9),
]
print('Cohort: N=1,524; deaths=115 (7.5%); de novo shock=171 (11.2%); composite=197 (12.9%)')
secondary = pd.DataFrame([
    {
        'Outcome': label,
        'Events': events,
        'Prevalence': f'{prevalence:.1f}%',
        'AUC': f\"{values['mean']:.3f}\",
        'SD': f\"{values['std']:.3f}\",
        '95% CI': f\"{values['ci_lo']:.3f}-{values['ci_hi']:.3f}\",
    }
    for label, values, events, prevalence in outcome_rows
])
display(secondary)
print('AUPRC: mortality=0.301; de novo shock=0.500; composite=0.635')
print('Top-3 feature importance:')
print(f\"  Mortality: eGFR ({three['mortality']['feature_importance']['egfr_igd']:.3f}), ureum ({three['mortality']['feature_importance']['ureum_igd']:.3f}), LVOT VTI ({three['mortality']['feature_importance']['lvot_vti_igd']:.3f})\")
print(f\"  De novo shock: LVOT VTI ({three['shock']['feature_importance']['lvot_vti_igd']:.3f}), SBP ({three['shock']['feature_importance']['sbp']:.3f}), APTT ({three['shock']['feature_importance']['aptt_value']:.3f})\")
print(f\"  Composite: LVOT VTI ({three['composite']['feature_importance']['lvot_vti_igd']:.3f}), SBP ({three['composite']['feature_importance']['sbp']:.3f}), eGFR ({three['composite']['feature_importance']['egfr_igd']:.3f})\")"""
    narrative = (
        "Mortality had the strongest discrimination. For the secondary outcomes, AUPRC was "
        "0.500 for de novo shock and 0.635 for the composite, compared with prevalence baselines "
        "of 0.112 and 0.129. Feature-importance rankings were outcome-specific and describe model "
        "usage rather than causal effects."
    )
    new_cells = [
        nbformat.v4.new_markdown_cell("## 7. Secondary outcomes: de novo cardiogenic shock and composite"),
        nbformat.v4.new_code_cell(code),
        nbformat.v4.new_markdown_cell(narrative),
    ]
    nb.cells[insertion:insertion] = new_cells
    nbformat.write(nb, SOURCE_NB)


if __name__ == "__main__":
    finalize_docx()
    finalize_notebook()
    print("Synchronized DOCX and source notebook.")
