#!/usr/bin/env python3
"""Complete TRIPOD+AI map and repair final reference-list details."""

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from pathlib import Path


docx = Path(__file__).resolve().parents[1] / "TESIS_FINAL.docx"
doc = Document(docx)


def set_paragraph(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


for p in doc.paragraphs:
    text = p.text
    text = text.replace("pp.3720-3361", "pp.3720-3826")
    text = text.replace(
        "*European Heart Journal  sampai  Cardiovascular Imaging*",
        "*European Heart Journal: Cardiovascular Imaging*",
    )
    text = text.replace("Atherosclerosis,No Longer", "Atherosclerosis: No Longer")
    text = text.replace("Syndromes,Diagnostic", "Syndromes: Diagnostic")
    if text != p.text:
        set_paragraph(p, text)

# Add the cited TRIPOD+AI paper in alphabetical position if absent.
if not any("bmj-2023-078378" in p.text for p in doc.paragraphs):
    anchor = next(p for p in doc.paragraphs if p.text.startswith("Chioncel, O. et al., 2020."))
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    inserted = Paragraph(new_p, anchor._parent)
    inserted.style = anchor.style
    inserted.add_run(
        "Collins, G.S. et al., 2024. TRIPOD+AI statement: updated guidance for reporting "
        "clinical prediction models that use regression or machine learning methods. "
        "*BMJ*, 385, p.e078378. https://doi.org/10.1136/bmj-2023-078378."
    )

# Convert the earlier 17-domain summary into an explicit map of all 27 TRIPOD+AI main items.
tripod = next(t for t in doc.tables if t.cell(0, 0).text == "Item" and "Lokasi" in t.cell(0, 1).text)
items = [
    ("1. Judul", "Judul dan abstrak menyebut model prediksi, luaran, populasi, dan validasi internal."),
    ("2. Abstrak", "Abstrak terstruktur memuat desain, N, kejadian, prediktor, validasi, dan performa."),
    ("3. Latar belakang", "Bagian 1.1 menjelaskan konteks klinis dan kebutuhan model."),
    ("4. Tujuan", "Bagian 1.3 memuat tujuan pengembangan dan validasi internal."),
    ("5. Sumber data", "Bagian 2.2 sampai 2.3: RME rumah sakit dan periode 2024 sampai 2025."),
    ("6. Partisipan", "Bagian 2.3 sampai 2.4: populasi, inklusi, eksklusi, dan alur seleksi."),
    ("7. Luaran", "Bagian 2.6 dan Tabel 2.1: definisi, waktu, sumber, dan pengodean mortalitas."),
    ("8. Prediktor", "Bagian 2.5 dan 2.7: definisi, satuan, sumber, dan waktu 13 prediktor."),
    ("9. Ukuran sampel", "Bagian 2.3 dan 4.10: N=1.524, 115 kejadian, serta keterbatasannya."),
    ("10. Data hilang", "Bagian 2.8 sampai 2.9: imputasi median dipelajari pada setiap lipatan pelatihan."),
    ("11. Pengembangan model", "Bagian 2.9: algoritma, hiperparameter, prapemrosesan, dan fitting."),
    ("12. Validasi internal", "Bagian 2.9: 10 seed, 5-fold StratifiedKFold, dan prediksi out-of-fold."),
    ("13. Ukuran performa", "Bagian 2.9 dan 3.2: AUC, AUPRC, Brier, sensitivitas, dan spesifisitas."),
    ("14. Pembentukan kelompok risiko", "Bagian 2.9 dan 3.8: definisi ambang keselamatan dan Youden."),
    ("15. Analisis tambahan", "Bagian 3.5 sampai 3.7: DCA, luaran sekunder, XGBoost, dan GRACE 2.0."),
    ("16. Perangkat analisis", "Bagian 2.9 dan berkas analisis menyebut algoritma, seed, serta pustaka."),
    ("17. Protokol dan registrasi", "Status protokol/registrasi tidak tersedia; keterbatasan pelaporan dinyatakan."),
    ("18. Sains terbuka", "Dokumentasi analisis tersedia; data individual dibatasi karena kerahasiaan pasien."),
    ("19. Keterlibatan pasien/publik", "Tidak dilakukan karena desain retrospektif berbasis RME."),
    ("20. Alur partisipan", "Bagian 3.1 dan Gambar 2.1 memuat seluruh tahap seleksi dan jumlah pasien."),
    ("21. Karakteristik partisipan", "Tabel 3.1 dan 3.2 memuat karakteristik kohort dan subkelompok."),
    ("22. Pengembangan model", "Bagian 3.2 sampai 3.5 memuat hasil model, fitur, ambang, dan kalibrasi."),
    ("23. Spesifikasi model", "Bagian 2.9 dan berkas analisis memuat 13 fitur dan hiperparameter lengkap."),
    ("24. Performa model", "Bagian 3.2 sampai 3.7 memuat estimasi, ketidakpastian, dan pembanding."),
    ("25. Keterbatasan", "Bagian 4.10 membahas bias, data hilang, pusat tunggal, dan validasi eksternal."),
    ("26. Interpretasi", "Bagian 4.1 sampai 4.9 menafsirkan hasil tanpa mengklaim kesiapan klinis."),
    ("27. Implikasi", "Bagian 5.2 memuat validasi eksternal, studi prospektif, dan studi dampak."),
]
while len(tripod.rows) > 1:
    tripod._tbl.remove(tripod.rows[-1]._tr)
for label, location in items:
    cells = tripod.add_row().cells
    cells[0].text = label
    cells[1].text = location

for p in doc.paragraphs:
    if p.text.startswith("Peta ini menunjukkan lokasi pelaporan 17 domain utama"):
        set_paragraph(
            p,
            "Checklist ini memetakan 27 item utama TRIPOD+AI ke lokasi pelaporan dalam tesis. "
            "Pemenuhan pelaporan tidak berarti validasi eksternal telah dilakukan.",
        )

doc.save(docx)
print(docx)
