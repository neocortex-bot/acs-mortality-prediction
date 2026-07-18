from docx import Document


path = "TESIS_FINAL.docx"
doc = Document(path)

updates = {
    159: "Ambang safety 0,018455 dipilih untuk membatasi false negative pada prediksi out-of-fold rerata. Pada ambang ini, sensitivitas adalah 98,3% (FN=2), spesifisitas 26,2%, PPV 9,8%, dan NPV 99,5%. Ambang Youden 0,103981 menghasilkan sensitivitas 71,3%, spesifisitas 82,0%, PPV 24,4%, dan NPV 97,2%.",
    160: "Ambang Youden 0,103981 memaksimalkan sensitivitas ditambah spesifisitas dikurangi satu pada kohort ini. Ambang tersebut mengidentifikasi 82 dari 115 kematian dan mengklasifikasikan benar 1.155 dari 1.409 pasien hidup. Karena ambang diturunkan dan dinilai pada kohort yang sama menggunakan prediksi out-of-fold, performanya memerlukan validasi eksternal.",
    163: "Matriks konfusi ambang safety disusun dari prediksi out-of-fold rerata seluruh 1.524 pasien. Model mengidentifikasi 113 dari 115 kematian (sensitivitas 98,3%) dan 369 dari 1.409 pasien hidup (spesifisitas 26,2%). Terdapat 2 false negative dan 1.040 false positive; total pasien yang terflag adalah 1.153. Hasil ini menggambarkan trade-off internal dan bukan jaminan keamanan klinis.",
}
for index, text in updates.items():
    paragraph = doc.paragraphs[index]
    for run in paragraph.runs:
        run.text = ""
    (paragraph.runs[0] if paragraph.runs else paragraph.add_run()).text = text

doc.save(path)
