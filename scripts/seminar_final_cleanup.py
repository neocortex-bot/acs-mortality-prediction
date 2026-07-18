#!/usr/bin/env python3
"""Final language, terminology, and consistency cleanup for TESIS_FINAL.docx."""

from pathlib import Path
from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "TESIS_FINAL.docx"


def set_text(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


doc = Document(DOCX)

# Replace complete paragraphs whose legacy wording is inaccurate or informal.
replacements = {
    148: (
        "Pada sebagian pasien, data ekokardiografi saat admisi tidak tercatat dalam RME karena "
        "ekokardiografi di sisi tempat tidur belum dilakukan ketika pasien diterima di IGD. "
        "Variasi praktik klinis di IGD, termasuk belum rutinnya POCUS ekokardiografi pada awal "
        "admisi dan perbedaan kelengkapan dokumentasi SOAP, berkontribusi terhadap "
        "ketidaktersediaan parameter ekokardiografi yang diperlukan untuk analisis. Kondisi ini "
        "merupakan salah satu alasan eksklusi pada tahap kendali mutu data."
    ),
    151: (
        "Tabel 3.1 menunjukkan bahwa pasien yang meninggal berusia lebih tua dan memiliki kelas "
        "Killip yang lebih berat saat admisi. Proporsi STEMI pada kelompok meninggal justru lebih "
        "rendah daripada kelompok hidup (62,6% dibandingkan 69,2%), dan perbedaannya tidak "
        "bermakna secara statistik (p=0,1736). Parameter hemodinamik dan laboratorium juga "
        "menunjukkan perbedaan antara kedua kelompok."
    ),
    161: (
        "Model Random Forest menunjukkan diskriminasi yang baik. Rerata AUC dari sepuluh "
        "pengulangan validasi silang adalah 0,8157 dengan simpangan baku 0,0075 dan rentang "
        "0,8024 sampai 0,8247. AUC vektor prediksi out-of-fold yang dirata-ratakan dari sepuluh "
        "pengulangan adalah 0,8189."
    ),
    255: (
        "Ukuran data yang sedang, yaitu 1.524 pasien dengan 115 kematian, mendukung pemilihan "
        "model yang relatif hemat parameter. Random Forest dapat menangkap interaksi tanpa "
        "transformasi eksplisit dan menunjukkan performa internal yang lebih baik daripada "
        "XGBoost. Namun, pemilihan algoritma tetap dapat dipengaruhi oleh ruang hiperparameter "
        "dan prosedur penalaan; tanpa validasi silang tersarang, selisih performa harus "
        "ditafsirkan secara hati-hati."
    ),
    309: (
        "1. Tiga belas prediktor dalam model mortalitas in-hospital mencakup fungsi ginjal "
        "(eGFR dan ureum), parameter ekokardiografi (LVOT VTI, LVEF, dan TAPSE), parameter darah "
        "(hemoglobin, APTT, dan kalium), usia, tekanan darah sistolik, denyut jantung, laju napas, "
        "dan kelas Killip. Berdasarkan Gini importance, eGFR dan ureum merupakan dua prediktor "
        "paling dominan. Kepentingan fitur menunjukkan kontribusi prediktif, bukan hubungan "
        "kausal atau kemaknaan statistik individual."
    ),
    313: (
        "5. Performa model untuk mortalitas (AUC 0,819) lebih baik daripada untuk syok "
        "kardiogenik baru (AUC 0,670) dan luaran komposit (AUC 0,742). Analisis luaran sekunder "
        "bersifat eksploratif dan tidak mengubah fokus utama penelitian."
    ),
    14: (
        "Sindrom Koroner Akut (SKA) merupakan manifestasi klinis kritis penyakit jantung koroner "
        "dan tetap menjadi penyebab utama kematian akibat penyakit kardiovaskular di seluruh dunia. "
        "SKA mencakup spektrum angina tidak stabil, infark miokard dengan elevasi segmen ST (STEMI), "
        "dan infark miokard tanpa elevasi segmen ST (NSTEMI), yang dipicu oleh iskemia miokard akut "
        "akibat penurunan mendadak aliran darah koroner (Thygesen et al., 2018). Data European "
        "Society of Cardiology menunjukkan bahwa penyakit kardiovaskular tetap menjadi penyebab "
        "kematian utama di dunia dan SKA berkontribusi besar terhadap morbiditas, mortalitas, dan "
        "biaya pelayanan kesehatan (Timmis et al., 2022). Di Indonesia, beban SKA meningkat karena "
        "tingginya faktor risiko tradisional dan keterbatasan akses terhadap pelayanan kesehatan."
    ),
    176: (
        "Matriks konfusi pada ambang keselamatan disusun dari prediksi out-of-fold rerata seluruh "
        "1.524 pasien. Model mengidentifikasi 113 dari 115 kematian (sensitivitas 98,3%) dan 369 "
        "dari 1.409 pasien hidup (spesifisitas 26,2%). Terdapat 2 negatif palsu dan 1.040 positif "
        "palsu; total pasien yang diklasifikasikan berisiko adalah 1.153. Hasil ini menggambarkan "
        "kompromi internal antara sensitivitas dan spesifisitas, bukan jaminan keamanan klinis."
    ),
    194: (
        "Mengingat ketidakseimbangan kelas dengan prevalensi mortalitas 7,5%, kurva presisi-recall "
        "memberikan informasi tambahan terhadap kurva ROC. AUPRC model adalah 0,301, lebih tinggi "
        "daripada baseline tanpa keterampilan sebesar 0,075."
    ),
    220: (
        "Probabilitas prediksi Random Forest digunakan untuk membentuk tiga strata risiko internal: "
        "risiko rendah dengan probabilitas <0,018455, risiko sedang dengan probabilitas 0,018455 "
        "sampai <0,103981, dan risiko tinggi dengan probabilitas >=0,103981. Label ward, HCU, dan "
        "ICU digunakan sebagai label konseptual dan bukan rekomendasi penempatan klinis otomatis."
    ),
    281: (
        "Luaran komposit (mortalitas dan/atau SKG) menunjukkan AUC 0,742, yang berada di antara "
        "performa untuk mortalitas (0,819) dan SKG (0,670). Analisis luaran komposit bersifat "
        "eksploratif karena menggabungkan komponen dengan karakteristik prediktif dan implikasi "
        "klinis yang berbeda."
    ),
    303: (
        "Namun, beberapa aspek TRIPOD+AI masih perlu ditingkatkan: (1) validasi eksternal pada "
        "kohort independen belum dilakukan; (2) analisis sensitivitas terkait data hilang belum "
        "menyeluruh; dan (3) analisis keadilan model pada subkelompok jenis kelamin, usia, dan "
        "komorbiditas perlu dilaporkan lebih terperinci. Penelitian lanjutan direncanakan untuk "
        "mengatasi keterbatasan tersebut."
    ),
}
for idx, value in replacements.items():
    set_text(doc.paragraphs[idx], value)

# Formal Indonesian wording and consistent terminology in all body paragraphs.
wording = {
    "dikroscek": "diperiksa silang",
    "ketidaktepatan koding": "ketidaktepatan pengodean",
    "dataset": "set data",
    "setting klinis real-time": "lingkungan klinis waktu nyata",
    "alert otomatis": "peringatan otomatis",
    "fairness analysis": "analisis keadilan model",
    "events yang terbatas": "kejadian yang terbatas",
    "events per variable": "kejadian per variabel",
    "single-center": "pusat tunggal",
    "missing data": "data hilang",
    "missingness": "proporsi data hilang",
    "not completely at random": "tidak hilang sepenuhnya secara acak",
    "Effect of treatment paradox": "Paradoks efek terapi",
    "treatment paradox": "paradoks terapi",
    "class imbalance": "ketidakseimbangan kelas",
    "precision": "presisi",
    "deep learning": "pembelajaran mendalam",
    "feature importance": "kepentingan fitur",
    "threshold safety": "ambang keselamatan",
    "Ambang safety": "Ambang keselamatan",
    "false negative": "negatif palsu",
    "false positive": "positif palsu",
    "raw out-of-fold predictions": "prediksi out-of-fold mentah",
}
for paragraph in doc.paragraphs:
    for old, new in wording.items():
        if old in paragraph.text:
            set_text(paragraph, paragraph.text.replace(old, new))

# Tables: remove all em dashes and standardize labels/source statement.
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                text = paragraph.text.replace("—", ",").replace("–", " sampai ")
                text = text.replace("Kriteria Objektif / Koding", "Kriteria Objektif/Pengodean")
                text = text.replace("Waktu Training", "Waktu Pelatihan")
                text = text.replace("threshold safety", "ambang keselamatan")
                text = text.replace("Makassar ACS Registry, pusat studi", "RME rumah sakit, pusat studi")
                text = text.replace("training fold", "lipatan pelatihan")
                text = text.replace("preprocessing", "prapemrosesan")
                if text != paragraph.text:
                    set_text(paragraph, text)

# A few English technical terms are retained where they are standard, but the clinical label is
# consistently "kelas Killip" in Indonesian prose.
for paragraph in doc.paragraphs:
    text = paragraph.text.replace("Killip class", "kelas Killip")
    if text != paragraph.text:
        set_text(paragraph, text)

# The earlier finalization script inserted the TRIPOD appendix before the abstract. Move its page
# break, heading, introduction, and table to the true end of the document, before sectPr.
body = doc.element.body
children = list(body.iterchildren())
if (
    len(children) >= 4
    and children[0].tag == qn("w:p")
    and children[1].tag == qn("w:p")
    and "LAMPIRAN: PETA PELAPORAN TRIPOD+AI" in "".join(children[1].itertext())
    and children[2].tag == qn("w:p")
    and children[3].tag == qn("w:tbl")
):
    appendix_nodes = children[:4]
    for node in appendix_nodes:
        body.remove(node)
    # The file has an early section-properties node before the abstract and a final one. The
    # appendix belongs before the final node, not before the first section break.
    section_nodes = body.findall(qn("w:sectPr"))
    sect_pr = section_nodes[-1] if section_nodes else None
    insert_at = body.index(sect_pr) if sect_pr is not None else len(body)
    for node in appendix_nodes:
        body.insert(insert_at, node)
        insert_at += 1

doc.save(DOCX)
print(DOCX)
