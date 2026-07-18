from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SOURCE = Path("TESIS_FINAL.docx")
BACKUP = Path("TESIS_FINAL.pre_seminar_audit.docx")


def replace_paragraph(paragraph, text):
    # Retain paragraph properties while replacing inconsistent legacy prose.
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


if not BACKUP.exists():
    BACKUP.write_bytes(SOURCE.read_bytes())

doc = Document(SOURCE)

replacements = {
    2: "Latar belakang: Mortalitas in-hospital pada sindrom koroner akut memerlukan stratifikasi risiko dini. Tujuan: Mengembangkan dan memvalidasi internal model Random Forest untuk mortalitas in-hospital pada pasien STEMI dan NSTEMI. Metode: Kohort retrospektif mencakup 1.524 pasien Killip I sampai III dengan 115 kematian. Tiga belas prediktor dianalisis menggunakan StratifiedKFold lima lipat pada sepuluh seed tetap tanpa kebocoran data; imputasi median dipelajari hanya dari data pelatihan setiap lipatan. Hasil: Rerata AUC antar-seed adalah 0,8157 (dibulatkan 0,816), sedangkan AUC vektor prediksi out-of-fold rerata adalah 0,8189 (dibulatkan 0,819). Brier score adalah 0,061 dan AUPRC 0,301. Ambang safety 0,018455 menghasilkan sensitivitas 98,3%, spesifisitas 26,2%, dan dua false negative. Ambang Youden 0,103981 menghasilkan sensitivitas 71,3% dan spesifisitas 82,0%. AUC GRACE 2.0 adalah 0,777 dibandingkan 0,819 untuk Random Forest. Kesimpulan: Model memiliki diskriminasi yang baik pada validasi internal, tetapi ambang klinis dan sistem triase memerlukan validasi eksternal sebelum penerapan.",
    6: "Background: Early mortality risk stratification is needed in acute coronary syndrome. Objective: To develop and internally validate a Random Forest model for in-hospital mortality. Methods: This retrospective cohort included 1,524 Killip I to III patients and 115 deaths. Thirteen predictors were evaluated using five-fold StratifiedKFold cross-validation over ten fixed seeds without data leakage; median imputation was fitted only on each training fold. Results: The mean AUC across seeds was 0.8157 (reported as 0.816), while the AUC of the averaged out-of-fold prediction vector was 0.8189 (reported as 0.819). The Brier score was 0.061 and AUPRC was 0.301. At the 0.018455 safety threshold, sensitivity was 98.3%, specificity was 26.2%, and two deaths were missed. At the 0.103981 Youden threshold, sensitivity was 71.3% and specificity was 82.0%. GRACE 2.0 AUC was 0.777 versus 0.819 for Random Forest. Conclusion: Internal discrimination was good, but the clinical thresholds and triage system require external validation before use.",
    37: "Random Forest dipilih karena dapat menangkap hubungan non-linear dan interaksi prediktor, sesuai untuk data tabular berukuran sedang, serta menyediakan ringkasan kepentingan fitur. Model menggunakan 13 prediktor yang ditetapkan berdasarkan relevansi klinis, ketersediaan saat admisi, kelengkapan data, dan pembatasan kompleksitas terhadap 115 kejadian. Random Forest tidak menangani nilai hilang secara langsung pada implementasi ini; karena itu, imputasi median dipelajari hanya dari data pelatihan setiap lipatan. XGBoost digunakan sebagai pembanding algoritmik pada analisis sekunder.",
    38: "Pusat Jantung Terpadu RSUP Dr. Wahidin Sudirohusodo menangani volume kasus SKA yang tinggi dengan spektrum keparahan yang luas. Belum tersedia model prediksi mortalitas in-hospital yang dikembangkan dan divalidasi secara khusus pada populasi pusat ini. Penelitian ini bertujuan mengembangkan dan melakukan validasi internal model Random Forest untuk mortalitas in-hospital pada pasien STEMI dan NSTEMI, menggunakan data yang tersedia dalam 24 jam pertama admisi.",
    59: "Penelitian ini merupakan studi analitik observasional dengan desain kohort retrospektif untuk mengembangkan dan melakukan validasi internal model prediksi mortalitas in-hospital pada pasien STEMI dan NSTEMI yang masuk melalui IGD.",
    115: "Variabel terikat utama adalah mortalitas in-hospital, yaitu kematian selama episode perawatan yang sama setelah admisi karena STEMI atau NSTEMI. Kejadian syok kardiogenik selama perawatan dicatat sebagai luaran sekunder dan tidak digunakan untuk menentukan luaran utama.",
    126: "Model utama adalah Random Forest dengan 500 pohon keputusan (n_estimators=500), kedalaman maksimum 6 (max_depth=6), dan minimal 5 sampel per daun (min_samples_leaf=5). Hiperparameter ditetapkan sebelum evaluasi multi-seed berdasarkan eksplorasi pengembangan; karena tidak digunakan nested cross-validation, performa dilaporkan sebagai validasi internal dan berpotensi optimistis. Evaluasi menggunakan StratifiedKFold lima lipat yang diulang pada sepuluh seed tetap. Pada setiap lipatan, median setiap prediktor dipelajari dari data pelatihan lalu diterapkan pada data validasi, sehingga tidak terjadi kebocoran informasi.",
    173: "AUPRC model adalah 0,301, lebih tinggi daripada prevalensi mortalitas 7,5%. Perbandingan ini bersifat deskriptif dan menunjukkan peningkatan pemeringkatan kelas minoritas dibandingkan baseline tanpa keterampilan; ketepatan klinis tetap bergantung pada ambang yang dipilih.",
    185: "Tiga fitur dengan Gini importance tertinggi adalah eGFR, ureum, dan LVOT VTI. Temuan ini menunjukkan bahwa fungsi ginjal dan parameter hemodinamik ekokardiografi berkontribusi pada prediksi model. Gini importance tidak menyatakan hubungan kausal dan dapat dipengaruhi oleh korelasi antarprediktor.",
    194: "Sistem triase menunjukkan gradien risiko: 0,5% pada kelompok ward, 3,8% pada kelompok HCU, dan 24,4% pada kelompok ICU. Kelompok ICU terdiri dari 336 pasien dan mencakup 82 dari 115 kematian. Ketiga strata diturunkan dan dinilai pada prediksi out-of-fold kohort yang sama; hasil ini merupakan validasi internal, bukan bukti keamanan alokasi tempat perawatan.",
    201: "Model menunjukkan AUC 0,819 untuk mortalitas, 0,742 untuk luaran komposit, dan 0,670 untuk syok kardiogenik baru. Analisis luaran sekunder bersifat eksploratif dan tidak mengubah fokus utama penelitian pada mortalitas in-hospital.",
    220: "Model Random Forest menunjukkan diskriminasi yang baik pada validasi internal. Rerata AUC antar-seed adalah 0,8157 dengan simpangan baku 0,0075, IK 95% 0,8110 sampai 0,8204, dan rentang 0,8024 sampai 0,8247. AUC vektor prediksi out-of-fold yang dirata-ratakan pada sepuluh seed adalah 0,8189. Brier score 0,061 menilai akurasi probabilitas secara keseluruhan, tetapi tidak menggantikan pemeriksaan calibration-in-the-large dan calibration slope. Pada ambang Youden 0,103981, sensitivitas adalah 71,3%, spesifisitas 82,0%, PPV 24,4%, dan NPV 97,2%. Pada ambang safety 0,018455, sensitivitas adalah 98,3%, spesifisitas 26,2%, dan terdapat dua false negative.",
    222: "Pengulangan validasi silang pada sepuluh seed menilai variasi akibat pembagian lipatan. Rerata AUC antar-seed adalah 0,8157 dengan simpangan baku 0,0075 dan rentang 0,8024 sampai 0,8247. Variasi ini menunjukkan bahwa estimasi performa tetap bergantung pada resampling; validasi eksternal diperlukan untuk menilai transportabilitas.",
    223: "Brier score 0,061 menunjukkan galat kuadrat probabilitas yang rendah dalam kohort ini, tetapi nilainya dipengaruhi oleh prevalensi mortalitas 7,5%. Kurva kalibrasi memberikan pemeriksaan visual. Calibration-in-the-large dan calibration slope belum dilaporkan, sehingga klaim kalibrasi tidak boleh melampaui bukti yang tersedia.",
    225: "Pada analisis sekunder, rerata AUC Random Forest adalah 0,816 dan AUC XGBoost adalah 0,789. Perbandingan ini bersifat internal dan tidak membuktikan keunggulan umum Random Forest pada populasi lain.",
    226: "Ukuran dataset yang sedang, 1.524 pasien dengan 115 kematian, mendukung pemilihan model yang relatif parsimonious. Random Forest menangkap interaksi tanpa transformasi eksplisit dan memberikan performa internal yang lebih baik daripada pembanding XGBoost. Namun, pemilihan algoritma tetap dapat dipengaruhi oleh ruang hiperparameter dan prosedur tuning; tanpa nested cross-validation, selisih performa harus ditafsirkan secara hati-hati.",
    234: "Kelompok risiko rendah atau ward, dengan probabilitas kurang dari 0,018455, terdiri dari 371 pasien dan mencakup 2 kematian (0,5%). NPV pada ambang ini adalah 99,5%. Kelompok ini tidak dapat disebut aman untuk ward berdasarkan validasi internal saja; keputusan lokasi perawatan tetap memerlukan penilaian klinis dan validasi prospektif eksternal.",
    235: "Kelompok risiko sedang atau HCU, dengan probabilitas 0,018455 sampai kurang dari 0,103981, mencakup 817 pasien dengan 31 kematian (3,8%). Kelompok risiko tinggi atau ICU, dengan probabilitas sekurang-kurangnya 0,103981, mencakup 336 pasien dengan 82 kematian (24,4%). Nama ward, HCU, dan ICU adalah usulan strata triase, bukan rekomendasi penempatan otomatis.",
    241: "Pada populasi yang sama, AUC GRACE 2.0 adalah 0,777 dan AUC Random Forest adalah 0,819, dengan selisih AUC 0,042 (bootstrap p=0,029; IK 95% 0,003 sampai 0,084). Pada perbandingan klasifikasi berpasangan di ambang risiko 20%, uji McNemar menghasilkan p<0,001. Hasil diskriminasi dan hasil McNemar menjawab pertanyaan yang berbeda, sehingga keduanya dilaporkan terpisah.",
    247: "Model menggunakan 13 prediktor yang tersedia dalam 24 jam pertama dan menyediakan feature importance untuk interpretasi global. Kebutuhan LVOT VTI dan TAPSE dapat membatasi penerapan di fasilitas tanpa ekokardiografi dini. Model belum siap untuk penggunaan klinis karena belum menjalani validasi eksternal, evaluasi prospektif, dan studi dampak implementasi.",
    281: "2. Rerata AUC Random Forest antar-seed adalah 0,8157 (dibulatkan 0,816), dengan simpangan baku 0,0075 dan IK 95% 0,8110 sampai 0,8204. AUC vektor prediksi out-of-fold rerata adalah 0,8189 (dibulatkan 0,819), sedangkan Brier score adalah 0,061. Pada ambang Youden 0,103981, sensitivitas adalah 71,3% dan spesifisitas 82,0%. Pada ambang safety 0,018455, sensitivitas adalah 98,3%, spesifisitas 26,2%, dan terdapat dua false negative.",
    282: "3. Sistem triase internal membagi 1.524 pasien menjadi ward sebanyak 371 pasien dengan 2 kematian (0,5%), HCU sebanyak 817 pasien dengan 31 kematian (3,8%), dan ICU sebanyak 336 pasien dengan 82 kematian (24,4%). Strata ini belum boleh digunakan sebagai aturan penempatan klinis sebelum validasi eksternal dan evaluasi prospektif.",
    283: "4. Pada populasi yang sama, AUC Random Forest adalah 0,819 dan AUC GRACE 2.0 adalah 0,777. Selisih AUC adalah 0,042 (bootstrap p=0,029), sedangkan perbandingan klasifikasi berpasangan pada ambang risiko 20% menghasilkan McNemar p<0,001. Perbandingan dengan XGBoost bersifat sekunder, dengan AUC 0,789.",
}

for index, text in replacements.items():
    replace_paragraph(doc.paragraphs[index], text)

# Correct the duplicated figure number and its in-text reference.
replace_paragraph(doc.paragraphs[214], "Gambar 3.12 Stratifikasi Risiko untuk Prediksi Mortalitas Berdasarkan Model Random Forest")
replace_paragraph(doc.paragraphs[257], doc.paragraphs[257].text.replace("Gambar 3.11", "Gambar 3.12"))

# Main outcome in the operational definition table.
t0 = doc.tables[0]
t0.cell(2, 0).text = "Mortalitas in-hospital"
t0.cell(2, 1).text = "Kematian selama episode perawatan rumah sakit yang sama setelah admisi dengan STEMI atau NSTEMI"
t0.cell(2, 2).text = "Biner"
t0.cell(2, 3).text = "Meninggal=1; hidup saat pulang=0, berdasarkan status keluar rumah sakit"

# Exact corrected feature importance and model summary values.
doc.tables[3].cell(3, 2).text = "0,0985"
doc.tables[6].cell(7, 1).text = "13"
doc.tables[6].cell(7, 2).text = "13"
doc.tables[7].cell(6, 3).text = "p<0,001"

# Replace dash glyphs throughout body and tables. Hyphens in compounds are retained.
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        run.text = run.text.replace("—", ",").replace("–", " sampai ")
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.text = run.text.replace("—", ",").replace("–", " sampai ")

# Add an explicit 17-item reporting map so seminar reviewers can locate every item.
doc.add_page_break()
p = doc.add_paragraph("LAMPIRAN: PETA PELAPORAN TRIPOD+AI")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].bold = True
doc.add_paragraph(
    "Peta ini menunjukkan lokasi pelaporan 17 domain utama. Pemenuhan pelaporan tidak berarti validasi eksternal telah dilakukan."
)
items = [
    ("1. Judul dan jenis studi", "Judul/abstrak: model prediksi, luaran, populasi, dan validasi internal disebutkan."),
    ("2. Ringkasan terstruktur", "Abstrak: desain, ukuran sampel, kejadian, prediktor, validasi, dan metrik utama."),
    ("3. Latar belakang dan tujuan", "Bagian 1.1 sampai 1.3: konteks klinis, model pembanding, dan tujuan."),
    ("4. Sumber data dan waktu", "Bagian 2.2 sampai 2.3: Makassar ACS Registry, pusat studi, dan periode 2024 sampai 2025."),
    ("5. Partisipan", "Bagian 2.3 sampai 2.4: populasi, total sampling, inklusi, dan eksklusi."),
    ("6. Luaran", "Bagian 2.6 dan Tabel 2.1: mortalitas in-hospital, sumber, waktu, dan pengodean."),
    ("7. Prediktor", "Bagian 2.5 dan 2.7: definisi 13 prediktor, satuan, dan waktu pengukuran."),
    ("8. Ukuran sampel", "Bagian 2.3 dan 4.10: N=1.524, 115 kejadian, dan keterbatasan event count."),
    ("9. Data hilang", "Bagian 2.8 sampai 2.9: median per-fold dipelajari hanya dari training fold."),
    ("10. Pengembangan model", "Bagian 2.9: algoritma, hiperparameter, serta prosedur fitting."),
    ("11. Validasi internal", "Bagian 2.9: 10 seed kali 5-fold StratifiedKFold dan prediksi out-of-fold."),
    ("12. Ukuran performa", "Bagian 2.9 dan 3.2: AUC, AUPRC, Brier score, sensitivitas, spesifisitas, PPV, dan NPV."),
    ("13. Alur dan karakteristik partisipan", "Bagian 3.1, Gambar 3.1, serta Tabel 3.1 sampai 3.2."),
    ("14. Spesifikasi model", "Bagian 2.9 dan berkas analisis: 13 fitur, preprocessing, hiperparameter, dan seed."),
    ("15. Hasil performa dan ketidakpastian", "Bagian 3.2 sampai 3.7: rerata, simpangan baku, IK, rentang, ambang, dan pembanding GRACE."),
    ("16. Keterbatasan dan interpretasi", "Bagian 4.9 sampai 4.10: single-center, retrospektif, bias seleksi, missingness, dan tanpa validasi eksternal."),
    ("17. Ketersediaan dan rencana validasi", "Bagian 5.2 dan dokumentasi analisis: reproduksi lokal, pembatasan akses data pasien, validasi eksternal prospektif, dan studi implementasi."),
]
table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
table.rows[0].cells[0].text = "Item"
table.rows[0].cells[1].text = "Lokasi dan status pelaporan"
for cell in table.rows[0].cells:
    shade(cell, "D9EAF7")
    for run in cell.paragraphs[0].runs:
        run.bold = True
for label, location in items:
    cells = table.add_row().cells
    cells[0].text = label
    cells[1].text = location
for row in table.rows:
    for cell in row.cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

# Standardize existing table header formatting and repeat headers across pages.
for table in doc.tables:
    if not table.rows:
        continue
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)
    for cell in table.rows[0].cells:
        shade(cell, "D9EAF7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True

doc.save(SOURCE)
print(f"Updated {SOURCE}; backup at {BACKUP}")
