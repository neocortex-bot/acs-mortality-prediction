#!/usr/bin/env python3
"""
Fix DOCX with correct churn data:
1. Replace STROBE flowchart (rId30) with v2
2. Update narrative P144 with refined language
3. dr_TM/dr_NP as data entry officers
"""
from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import os, zipfile

DOCX = 'TESIS_FINAL.docx'
STROBE_V2 = 'figures/strobe_flowchart_v2.png'

doc = Document(DOCX)
body = doc.element.body

# ─── 1. REPLACE STROBE IMAGE (rId30 = image25.png) ───
# Find the paragraph with the old image
for i, p in enumerate(doc.paragraphs):
    xml = etree.fromstring(p._element.xml.encode())
    ns_a = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    blips = xml.findall(f'.//{{{ns_a}}}blip')
    for b in blips:
        embed = b.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
        if embed and embed in doc.part.rels:
            tgt = doc.part.rels[embed].target_ref
            if 'image25' in tgt:
                # Found the old STROBE paragraph
                # Remove it
                p._element.getparent().remove(p._element)
                print(f"  Removed old STROBE at P{i}")

# Insert new STROBE after "2.11 Alur Penelitian"
for i, p in enumerate(doc.paragraphs):
    if '2.11 Alur Penelitian' in p.text:
        # Create new paragraph with image
        new_para = OxmlElement('w:p')
        new_run = OxmlElement('w:r')
        # Use add_picture to create proper XML
        temp_p = doc.paragraphs[i]
        temp_run = temp_p.add_run()
        temp_run.add_picture(STROBE_V2, width=Cm(14))
        
        # Move this run to our new paragraph
        last_run = list(temp_p._element.iter(qn('w:r')))[-1]
        temp_p._element.remove(last_run)
        new_para.append(last_run)
        
        # Insert after "2.11 Alur Penelitian"
        temp_p._element.addnext(new_para)
        
        # Also update caption
        for j in range(i, i+5):
            if j < len(doc.paragraphs) and 'Diagram alur seleksi' in doc.paragraphs[j].text:
                # Update text
                for r in list(doc.paragraphs[j]._element.iter(qn('w:r'))):
                    for t in list(r.iter(qn('w:t'))):
                        t.text = 'Gambar 2.1 Diagram alur seleksi pasien penelitian'
                break
        
        print(f"  Inserted STROBE v2 after P{i}")
        break

# ─── 2. UPDATE NARRATIVE P144 ───
new_churn_text = (
    "Dari total 1.952 pasien yang tercatat dalam Makassar ACS Registry, sebanyak 379 pasien "
    "tidak memiliki data terstruktur (cleaned_data) karena diagnosis yang tidak memenuhi kriteria "
    "inklusi SKA — seperti sepsis, pneumonia, atau gagal jantung — data rekam medis yang tidak "
    "memadai untuk ekstraksi parameter prediktor, atau pasien yang meninggal sebelum sempat "
    "melengkapi dokumentasi awal. Sebagian dari kasus ini merupakan data yang di-entry oleh "
    "residen atau koasisten yang belum menyelesaikan pelatihan ekstraksi data terstruktur; variasi "
    "kompetensi dokumentasi ini menyebabkan sejumlah rekam medis tidak dapat diproses secara seragam "
    "ke dalam format analitik. Proses seleksi ini menghasilkan 1.573 pasien dengan data terstruktur "
    "yang lengkap."
    "\n\n"
    "Dari 1.573 pasien tersebut, 48 pasien dikeluarkan melalui mekanisme validasi internal "
    "(pat_exclude=True) karena terdeteksi memiliki data yang tidak memadai untuk analisis — meliputi "
    "parameter ekokardiografi dan laboratorium yang tidak terekstrak secara utuh, temuan halusinasi "
    "pada data terstruktur, atau pasien yang pulang paksa sebelum outcome dapat dievaluasi. "
    "Lima belas dari 48 pasien ini berada dalam Killip IV (syok kardiogenik manifes saat presentasi "
    "awal), dan 15 pasien di antaranya meninggal. Setelah eksklusi validasi, diperoleh 1.525 pasien "
    "yang memenuhi syarat analisis. Satu pasien dengan Killip IV yang terlewat dari filter eksklusi "
    "juga tidak diikutsertakan karena model ini dirancang untuk memprediksi mortalitas pada pasien "
    "yang belum berada dalam keadaan syok — bersifat prediktif, bukan diagnostik terhadap kondisi "
    "yang sudah nyata."
    "\n\n"
    "Perlu dicatat bahwa dua orang data entry officer (dr_TM dan dr_NP) berkontribusi dalam "
    "ekstraksi data untuk 11 pasien dalam database ini. Tingkat missing data ekokardiografi pada "
    "data yang mereka ekstrak lebih tinggi — LVEF 36%, LVOT VTI 73%, dan TAPSE 36% — dibandingkan "
    "keseluruhan populasi. Hal ini mencerminkan variasi individual dalam konsistensi ekstraksi "
    "data dari rekam medis, dan tidak berkaitan dengan kualitas tatalaksana klinis oleh DPJP."
)

# Find P144
target_idx = None
for i, p in enumerate(doc.paragraphs):
    if 'total 1.573 pasien yang tercatat dalam registry' in p.text or 'total 1.573 pasien' in p.text:
        target_idx = i
        break

if target_idx is not None:
    p = doc.paragraphs[target_idx]
    p_elem = p._element
    
    # Remove existing runs
    for r in list(p_elem.iter(qn('w:r'))):
        p_elem.remove(r)
    
    # Split by double newline
    paragraphs = new_churn_text.strip().split('\n\n')
    
    # First paragraph: update in-place
    r1 = OxmlElement('w:r')
    t1 = OxmlElement('w:t')
    t1.text = paragraphs[0]
    t1.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r1.append(t1)
    p_elem.append(r1)
    
    # Insert remaining paragraphs after
    current = p_elem
    for para_text in paragraphs[1:]:
        new_p = OxmlElement('w:p')
        new_r = OxmlElement('w:r')
        new_t = OxmlElement('w:t')
        new_t.text = para_text
        new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        new_r.append(new_t)
        new_p.append(new_r)
        current.addnext(new_p)
        current = new_p
    
    print(f"  Narrative updated at P{target_idx} ({len(paragraphs)} paragraphs)")
else:
    print("  WARNING: Could not find P144 narrative")

# ─── SAVE ───
doc.save(DOCX)
print(f"✓ DOCX saved: {os.path.getsize(DOCX)} bytes")
