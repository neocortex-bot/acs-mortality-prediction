#!/usr/bin/env python3
"""Execute the notebook and fail when notebook, JSON, or DOCX metrics diverge."""
import json
import contextlib
import io
import subprocess
import sys
from pathlib import Path

import nbformat
from docx import Document

ROOT = Path(__file__).resolve().parent
EXECUTED = ROOT / "gm-acs-mortality-prediction-executed.ipynb"

def fail(messages):
    print("VERIFICATION FAILED")
    for message in messages: print(" -", message)
    raise SystemExit(1)

try:
    subprocess.run([
        sys.executable, "-m", "jupyter", "nbconvert", "--to", "notebook", "--execute",
        "gm-acs-mortality-prediction.ipynb", "--output", EXECUTED.name,
        "--ExecutePreprocessor.timeout=900",
    ], cwd=ROOT, check=True)
except subprocess.CalledProcessError:
    print("nbconvert kernel sockets are unavailable; using sequential in-process execution.")
    source_nb = nbformat.read(ROOT / "gm-acs-mortality-prediction.ipynb", as_version=4)
    namespace = {"__name__": "__main__"}
    old_cwd = Path.cwd()
    try:
        import os
        os.chdir(ROOT)
        execution_count = 0
        for cell in source_nb.cells:
            if cell.cell_type != "code":
                continue
            execution_count += 1
            buffer = io.StringIO()
            namespace["display"] = lambda obj, _b=buffer: print(obj, file=_b)
            with contextlib.redirect_stdout(buffer), contextlib.redirect_stderr(buffer):
                exec(compile(cell.source, f"notebook-cell-{execution_count}", "exec"), namespace)
            cell.execution_count = execution_count
            cell.outputs = [nbformat.v4.new_output("stream", name="stdout", text=buffer.getvalue())]
    finally:
        os.chdir(old_cwd)
    nbformat.write(source_nb, EXECUTED)

r = json.loads((ROOT / "validation_results.json").read_text())
three = json.loads((ROOT / "three_outcomes_results.json").read_text())
nb = nbformat.read(EXECUTED, as_version=4)
output_text = "\n".join(
    str(o.get("text", "")) + str(o.get("data", {}).get("text/plain", ""))
    for cell in nb.cells if cell.cell_type == "code" for o in cell.get("outputs", [])
)
doc = Document(ROOT / "TESIS_FINAL.docx")
doc_text = "\n".join(p.text for p in doc.paragraphs) + "\n" + "\n".join(
    cell.text for table in doc.tables for row in table.rows for cell in row.cells
)
m, t, g = r["metrics"], r["thresholds"], r["grace"]
checks = {
    "notebook cohort": (f'N={r["dataset"]["n"]:,}' in output_text),
    "notebook seed AUC": (f'{m["auc_mean"]:.4f}' in output_text),
    "notebook Brier": (f'{m["brier_mean"]:.4f}' in output_text),
    "notebook AUPRC": (f'{m["auprc_mean"]:.4f}' in output_text),
    "notebook safety threshold": (f'{t["safety"]:.6f}' in output_text),
    "notebook Youden threshold": (f'{t["youden"]:.6f}' in output_text),
    "notebook GRACE AUC": (f'{g["auc"]:.4f}' in output_text),
    "DOCX cohort": ("1.524" in doc_text and "115" in doc_text),
    "DOCX seed AUC": (f'{m["auc_mean"]:.4f}'.replace(".", ",") in doc_text),
    "DOCX Brier": (f'{m["brier_mean"]:.4f}'.replace(".", ",") in doc_text),
    "DOCX AUPRC": (f'{m["auprc_mean"]:.4f}'.replace(".", ",") in doc_text),
    "DOCX safety threshold": (f'{t["safety"]:.6f}'.replace(".", ",") in doc_text),
    "DOCX Youden threshold": (f'{t["youden"]:.6f}'.replace(".", ",") in doc_text),
    "DOCX GRACE AUC": (f'{g["auc"]:.3f}'.replace(".", ",") in doc_text),
    "notebook three-outcome AUCs": all(f'{three[k]["mean"]:.3f}' in output_text for k in ("mortality", "shock", "composite")),
    "notebook secondary counts": all(value in output_text for value in ("115", "171", "197")),
    "DOCX three-outcome AUCs": all(f'{three[k]["mean"]:.3f}'.replace(".", ",") in doc_text for k in ("mortality", "shock", "composite")),
    "DOCX secondary counts": all(value in doc_text for value in ("115", "171", "197")),
    "DOCX secondary AUPRC": ("0,500" in doc_text and "0,635" in doc_text),
    "DOCX secondary comparison table": all(value in doc_text for value in ("0,805-0,833", "0,736-0,757", "0,761-0,777")),
}
errors = [name for name, passed in checks.items() if not passed]
if errors: fail(errors)

print("VERIFICATION PASSED")
print(f'N={r["dataset"]["n"]}; deaths={r["dataset"]["deaths"]}')
print(f'Seed AUC={m["auc_mean"]:.4f} ± {m["auc_sd"]:.4f}; 95% CI {m["auc_ci95"][0]:.4f}-{m["auc_ci95"][1]:.4f}; range {m["auc_min"]:.4f}-{m["auc_max"]:.4f}')
print(f'Mean-OOF AUC={m["ensemble_auc"]:.4f}; Brier={m["brier_mean"]:.4f}; AUPRC={m["auprc_mean"]:.4f}')
print(f'Safety={t["safety"]:.6f}; sensitivity={t["safety_metrics"]["sensitivity"]:.1%}; specificity={t["safety_metrics"]["specificity"]:.1%}')
print(f'Youden={t["youden"]:.6f}; sensitivity={t["youden_metrics"]["sensitivity"]:.1%}; specificity={t["youden_metrics"]["specificity"]:.1%}')
print(f'GRACE AUC={g["auc"]:.4f}; McNemar p={g["threshold_20pct"]["mcnemar_p"]:.4g}')
